"""LangChain GAIA agent.

This file is the product-under-test. It can run by itself for manual tracing,
and eval workflows can import `solve` as the mvp-evals solver.
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import math
import os
import re
import time
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from pydantic import SecretStr

from evaris import Span, Trace

ZAI_BASE_URL = "https://api.z.ai/api/paas/v4"
DEFAULT_STUDIO_ENDPOINT = "http://localhost:3001/v1/traces"
AGENT_DIR = Path(__file__).resolve().parent
ROOT = AGENT_DIR.parents[1]

DEFAULT_CONFIG: dict[str, Any] = {
    "candidate_id": "baseline",
    "task_model": "glm-5",
    "max_output_tokens_per_call": 16384,
    "request_timeout_seconds": 90,
    "recursion_limit": 25,
    "max_iterations": 10,
}

DEFAULT_ARTIFACT: dict[str, Any] = {
    "system_prompt": (
        "You are a careful GAIA benchmark agent. Use tools when needed, verify factual "
        "answers, and return only the short final answer."
    ),
    "tool_policy": (
        "Use calculator for arithmetic. Use read_attachment when the task includes a file. "
        "Use web_search to look up facts, names, dates, or any information not in the prompt. "
        "Keep tool calls within budget. Prefer web_search over guessing."
    ),
    "answer_format_policy": "Return the final answer as a short string with no extra explanation.",
    "verification_policy": (
        "Before answering, verify that the final answer directly addresses the question."
    ),
}


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text())


def load_json_or_default(path: Path | None, default: dict[str, Any]) -> dict[str, Any]:
    if path is None or not path.exists():
        return dict(default)
    return default | load_json(path)


def artifact_hash(artifact: dict[str, Any]) -> str:
    payload = json.dumps(artifact, sort_keys=True)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:12]


def configure_environment() -> None:
    load_dotenv(ROOT / ".env")
    load_dotenv(AGENT_DIR / ".env")

    if not os.environ.get("OPENAI_API_KEY") and os.environ.get("ZAI_API_KEY"):
        os.environ["OPENAI_API_KEY"] = os.environ["ZAI_API_KEY"]

    os.environ.setdefault("OPENAI_BASE_URL", ZAI_BASE_URL)
    os.environ.setdefault("OTEL_SERVICE_NAME", "langchain-gaia-agent")
    os.environ.setdefault(
        "OTEL_EXPORTER_OTLP_TRACES_ENDPOINT",
        os.environ.get("ETRACE_STUDIO_OTLP_ENDPOINT", DEFAULT_STUDIO_ENDPOINT),
    )


def init_tracing() -> None:
    import etrace
    from etrace.otel import OtelExporter

    exporters = [etrace.InMemoryExporter()]  # Always keep in-memory for scorers
    if os.environ.get("ETRACE_EXPORT", "studio") == "studio":
        exporters.append(OtelExporter())

    etrace.init(
        service_name="langchain-gaia-agent",
        environment=os.environ.get("ETRACE_ENVIRONMENT", "local"),
        exporters=exporters,
        calculate_costs=True,
    )


def extract_text_from_agent_result(result: dict[str, Any]) -> str:
    messages = result.get("messages", [])
    for msg in reversed(messages):
        content = getattr(msg, "content", None)
        if isinstance(content, str) and content.strip():
            text = content.strip()
            # If multi-line, return only the last non-empty line (capped 200 chars)
            # so exact-match scoring sees the bare answer (e.g. "3")
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            if len(lines) > 1:
                return lines[-1][:200]
            return text[:200]
    return str(result)


def _build_iteration_limiter(max_iterations: int):
    """Build middleware that caps agent loop iterations.

    After max_iterations model calls, strips tool_calls from the AI message
    so the graph exits the loop instead of calling tools again.
    This prevents GraphRecursionError without inflating the recursion limit.
    """
    from langchain.agents.middleware import AgentMiddleware
    from langchain_core.messages import AIMessage

    class IterationLimiter(AgentMiddleware):
        def __init__(self, max_iters: int):
            super().__init__()
            self.max_iters = max_iters

        def before_model(self, state, runtime):
            messages = list(state.get("messages", []))
            ai_count = sum(1 for m in messages if isinstance(m, AIMessage))
            if ai_count >= self.max_iters:
                # Force-stop: replace the last AI message's tool_calls
                for msg in reversed(messages):
                    if isinstance(msg, AIMessage) and msg.tool_calls:
                        # Replace with a text-only response to break the loop
                        stripped = AIMessage(
                            content=msg.content or "I've reached my tool call limit. Returning my best answer based on the information gathered so far.",
                            id=msg.id,
                        )
                        messages[messages.index(msg)] = stripped
                        return {"messages": messages}
            return None

    return IterationLimiter(max_iterations)


def build_agent(config: dict[str, Any], artifact: dict[str, Any]):
    from langchain.agents import create_agent
    from langchain_core.messages import HumanMessage
    from langchain_core.tools import tool
    from langchain_openai import ChatOpenAI

    @tool
    def calculator(expression: str) -> str:
        """Evaluate a simple arithmetic expression. Use for math calculations."""
        if not re.fullmatch(r"[0-9\s+\-*/().%^]+", expression):
            return "calculator error: unsupported expression"
        try:
            expr = expression.replace("^", "**")
            value = eval(expr, {"__builtins__": {}}, {"sqrt": math.sqrt})  # noqa: S307
            return str(value)
        except Exception as exc:
            return f"calculator error: {exc}"

    @tool
    def read_attachment(path: str) -> str:
        """Read a local task attachment file when the GAIA task includes one.
        Supports .txt, .py, .xlsx, .docx, and plain text files."""
        file_path = Path(path)
        if not file_path.is_absolute():
            file_path = AGENT_DIR / file_path
        if not file_path.exists():
            return f"attachment not found: {file_path}"
        if file_path.stat().st_size > 500_000:
            return f"attachment too large for MVP reader: {file_path}"

        suffix = file_path.suffix.lower()

        # Excel — read cell values + fill colors
        if suffix == ".xlsx":
            try:
                import openpyxl

                wb = openpyxl.load_workbook(file_path, data_only=True)
                lines: list[str] = []
                for ws in wb.worksheets:
                    lines.append(f"[Sheet: {ws.title}]")
                    for row in ws.iter_rows():
                        cells: list[str] = []
                        for cell in row:
                            fill = ""
                            if (
                                cell.fill
                                and cell.fill.fgColor
                                and getattr(cell.fill.fgColor, "type", None)
                                != "theme"
                            ):
                                raw = str(cell.fill.fgColor.rgb or "")
                                # strip alpha prefix if present
                                if len(raw) == 8 and raw[0:2] == "FF":
                                    raw = raw[2:]
                                if raw.isalnum():
                                    fill = f" [fill=#{raw}]"
                            val = str(cell.value) if cell.value is not None else ""
                            cells.append(f"{cell.coordinate}:{val}{fill}")
                        lines.append(" | ".join(cells))
                return "\n".join(lines)
            except Exception as exc:
                return f"xlsx parse error: {exc}"

        # Word doc
        if suffix == ".docx":
            try:
                from docx import Document

                doc = Document(str(file_path))
                parts: list[str] = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                for table in doc.tables:
                    parts.append("[TABLE]")
                    for row in table.rows:
                        parts.append(" | ".join(cell.text for cell in row.cells))
                return "\n".join(parts)
            except Exception as exc:
                return f"docx parse error: {exc}"

        # Default: plain text (txt, py, csv, etc.)
        return file_path.read_text(errors="replace")

    @tool
    def web_search(query: str) -> str:
        """Search the web for information. Use for factual lookups, names, dates, and any information not in the prompt. Returns a summary of search results."""
        import urllib.parse

        encoded_query = urllib.parse.quote_plus(query)
        search_url = f"https://html.duckduckgo.com/html/?q={encoded_query}"
        try:
            resp = httpx.get(
                search_url,
                headers={"User-Agent": "Mozilla/5.0 (compatible; GAIA-agent/1.0)"},
                follow_redirects=True,
                timeout=15.0,
            )
            resp.raise_for_status()
            # Extract result snippets from DuckDuckGo HTML
            import re as _re

            results: list[str] = []
            for match in _re.findall(
                r"class=\"result__snippet\"[^>]*>(.*?)</a>",
                resp.text,
                _re.DOTALL,
            ):
                clean = _re.sub(r"<[^>]+>", "", match).strip()
                if clean:
                    results.append(clean)
            for match in _re.findall(
                r"class=\"result__a\"[^>]*>(.*?)</a>",
                resp.text,
                _re.DOTALL,
            ):
                clean = _re.sub(r"<[^>]+>", "", match).strip()
                if clean:
                    results.append(f"[URL] {clean}")
            if not results:
                # Fallback: return first 2000 chars of cleaned text
                clean = _re.sub(r"<script[^>]*>.*?</script>", "", resp.text, flags=_re.DOTALL)
                clean = _re.sub(r"<style[^>]*>.*?</style>", "", clean, flags=_re.DOTALL)
                clean = _re.sub(r"<[^>]+>", " ", clean)
                clean = _re.sub(r"\s+", " ", clean).strip()
                return clean[:2000]
            return "\n".join(results[:10])
        except Exception as exc:
            return f"web_search error: {exc}"

    system_prompt = "\n\n".join(
        [
            artifact["system_prompt"],
            "Tool policy:\n" + artifact["tool_policy"],
            "Answer format:\n" + artifact["answer_format_policy"],
            "Verification policy:\n" + artifact["verification_policy"],
        ]
    )

    model = ChatOpenAI(
        model=config.get("task_model", "glm-5"),
        api_key=SecretStr(os.environ["OPENAI_API_KEY"]),
        base_url=os.environ.get("OPENAI_BASE_URL", ZAI_BASE_URL),
        temperature=0,
        max_completion_tokens=config.get("max_output_tokens_per_call", 16384),
        max_retries=2,
        timeout=config.get("request_timeout_seconds", 90),
    )

    max_iterations = config.get("max_iterations", 10)
    iteration_limiter = _build_iteration_limiter(max_iterations)

    agent = create_agent(
        model=model,
        tools=[calculator, read_attachment, web_search],
        system_prompt=system_prompt,
        middleware=[iteration_limiter],
        name="langchain-gaia-agent",
    )

    return agent, HumanMessage


async def solve(
    sample_input: Any,
    *,
    config: dict[str, Any] | None = None,
    artifact: dict[str, Any] | None = None,
) -> Trace:
    """Run the GAIA agent and return an evaris Trace.

    `sample_input` is expected to be a dict with GAIA task fields:
    id, level, split, input, file_path.
    """
    import etrace

    config = config or dict(DEFAULT_CONFIG)
    artifact = artifact or dict(DEFAULT_ARTIFACT)

    task = dict(sample_input)
    artifact_id = artifact_hash(artifact)
    prompt = (
        f"GAIA Level {task.get('level', 1)} task.\n\n"
        f"Question: {task['input']}\n\n"
        f"Attachment path: {task.get('file_path') or 'none'}\n\n"
        "Return only the final short answer."
    )

    attrs = {
        "benchmark": "gaia",
        "task_id": task.get("id"),
        "task_level": task.get("level", 1),
        "split": task.get("split", "level1_train"),
        "candidate_id": config.get("candidate_id", "baseline"),
        "model": config.get("task_model", "glm-5"),
        "artifact_hash": artifact_id,
    }

    agent, HumanMessage = build_agent(config, artifact)
    start = time.time()

    max_iterations = config.get("max_iterations", 10)
    effective_recursion_limit = max_iterations * 5  # each iteration = ~5 recursion steps (model + tools + edges)

    # Capture the in-memory exporter to collect all spans (incl. LangChain tool calls)
    mem_exporter = None
    for exp in etrace._exporters:
        if hasattr(exp, "_spans") and hasattr(exp, "get_finished_spans"):
            mem_exporter = exp
            break
    if mem_exporter:
        mem_exporter.clear()
    pre_span_count = len(mem_exporter._spans) if mem_exporter else 0

    with etrace.trace("gaia.task", kind="agent", input=task, attributes=attrs) as root_span:
        result = await asyncio.to_thread(
            agent.invoke,
            {"messages": [HumanMessage(content=prompt)]},
            config={
                "callbacks": [etrace.langchain_handler()],
                "recursion_limit": effective_recursion_limit,
            },
        )
        answer = extract_text_from_agent_result(result)
        root_span.output = answer
        trace_id = getattr(root_span, "trace_id", None)

    duration_ms = (time.time() - start) * 1000

    # Build evaris Span tree from collected etrace spans
    root = Span(
        name="gaia.task",
        kind="agent",
        input=task,
        output=answer,
        duration_ms=duration_ms,
        attributes=attrs
        | {
            "trace_id": trace_id,
            "duration_ms": duration_ms,
        },
    )

    if mem_exporter:
        # Get spans created during this run (after pre_span_count)
        run_spans = mem_exporter._spans[pre_span_count:]
        # Convert all etrace spans to evaris Spans and build the tree
        span_map: dict[str, tuple[Span, str | None]] = {}
        for es in run_spans:
            s = Span(
                name=es.name,
                kind=es.kind.value if hasattr(es.kind, 'value') else str(es.kind),
                input=es.input,
                output=es.output,
                duration_ms=(es.duration_ns or 0) / 1_000_000,
                attributes=dict(es.attributes) if es.attributes else {},
            )
            span_map[es.span_id] = (s, es.parent_span_id)

        # Find the root span (no parent) and use it as our root
        for span_id, (s, pid) in span_map.items():
            if pid is None:
                s.input = task
                s.output = answer
                s.attributes.update(attrs)
                s.attributes["trace_id"] = trace_id
                s.attributes["duration_ms"] = duration_ms
                root = s
                break

        # Build the tree: attach each child to its parent
        visited: set[str] = set()
        for span_id, (child, parent_id) in span_map.items():
            if parent_id and parent_id in span_map and span_id not in visited:
                _attach_child(span_map, child, parent_id, visited)

        # Any remaining unattached spans go under root
        for span_id, (child, parent_id) in span_map.items():
            if span_id not in visited and child is not root:
                root.children.append(child)
                visited.add(span_id)

    return Trace(input=task["input"], output=answer, root=root)


def _attach_child(span_map: dict[str, tuple], child: Span, parent_id: str, visited: set[str]) -> None:
    """Attach a child span to its parent by walking up the tree."""
    if parent_id in visited:
        return
    parent = span_map.get(parent_id)
    if parent:
        visited.add(parent_id)
        parent[0].children.append(child)
    else:
        # Parent not in map — orphan, will be caught by fallback
        pass


def load_manifest(path: Path) -> list[dict[str, Any]]:
    tasks: list[dict[str, Any]] = []
    with path.open() as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            row = json.loads(line)
            tasks.append(
                {
                    "id": str(row.get("id") or row.get("task_id")),
                    "level": int(row.get("level") or row.get("Level") or 1),
                    "split": str(row.get("split") or "level1_train"),
                    "input": str(row.get("input") or row.get("Question")),
                    "expected": row.get("expected") or row.get("Final answer"),
                    "file_path": row.get("file_path") or row.get("file_name"),
                }
            )
    return tasks


def demo_task() -> dict[str, Any]:
    return {
        "id": "demo-level1-arithmetic",
        "level": 1,
        "split": "level1_train",
        "input": "A rectangle has side lengths 17 and 23. What is its area?",
        "expected": "391",
        "file_path": None,
    }


def choose_task(args: argparse.Namespace) -> dict[str, Any]:
    if args.question:
        return {
            "id": args.task_id or "manual-question",
            "level": args.level,
            "split": args.split,
            "input": args.question,
            "expected": args.expected,
            "file_path": args.file_path,
        }

    if args.manifest:
        tasks = load_manifest(Path(args.manifest))
        if not tasks:
            raise RuntimeError(f"No tasks found in manifest: {args.manifest}")
        if args.index >= len(tasks):
            raise RuntimeError(
                f"Manifest has {len(tasks)} tasks; index {args.index} is out of range"
            )
        return tasks[args.index]

    return demo_task()


def save_run_record(task: dict[str, Any], trace: Trace) -> Path:
    output_dir = AGENT_DIR / ".evaris" / "agent-runs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{int(time.time())}-{task['id']}.json"
    output_path.write_text(
        json.dumps(
            {
                "task": task,
                "answer": trace.output,
                "trace_id": trace.root.attributes.get("trace_id"),
                "duration_ms": trace.root.attributes.get("duration_ms"),
                "root_attributes": trace.root.attributes,
            },
            indent=2,
            default=str,
        )
    )
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the LangChain GAIA agent once.")
    parser.add_argument(
        "--manifest", help="Local GAIA JSONL manifest. Real GAIA content should stay local."
    )
    parser.add_argument("--index", type=int, default=0, help="Task index in --manifest.")
    parser.add_argument("--question", help="Run a manual question instead of a manifest row.")
    parser.add_argument("--expected", help="Optional expected answer for local bookkeeping.")
    parser.add_argument("--task-id", help="Optional task id for --question.")
    parser.add_argument("--level", type=int, default=1)
    parser.add_argument("--split", default="level1_train")
    parser.add_argument("--file-path", default=None)
    parser.add_argument(
        "--config",
        default=None,
        help="Optional config JSON. Defaults to embedded config.",
    )
    parser.add_argument(
        "--artifact",
        default=None,
        help="Optional prompt artifact JSON. Defaults to embedded prompt.",
    )
    parser.add_argument(
        "--no-studio",
        action="store_true",
        help="Use etrace in-memory tracing instead of exporting to Studio.",
    )
    return parser.parse_args()


async def main() -> None:
    args = parse_args()
    configure_environment()
    if args.no_studio:
        os.environ["ETRACE_EXPORT"] = "memory"

    if not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("Set ZAI_API_KEY or OPENAI_API_KEY before running the agent.")

    init_tracing()
    config_path = Path(args.config) if args.config else None
    artifact_path = Path(args.artifact) if args.artifact else None
    config = load_json_or_default(config_path, DEFAULT_CONFIG)
    artifact = load_json_or_default(artifact_path, DEFAULT_ARTIFACT)
    task = choose_task(args)

    trace = await solve(task, config=config, artifact=artifact)
    output_path = save_run_record(task, trace)

    print(
        json.dumps(
            {
                "task_id": task["id"],
                "answer": trace.output,
                "trace_id": trace.root.attributes.get("trace_id"),
                "studio_endpoint": os.environ.get("OTEL_EXPORTER_OTLP_TRACES_ENDPOINT"),
                "saved": str(output_path),
            },
            indent=2,
            default=str,
        )
    )


if __name__ == "__main__":
    asyncio.run(main())
